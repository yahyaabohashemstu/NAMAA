from flask import Flask, request, send_file, jsonify
from flask_cors import CORS
from docx import Document
import requests
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches
from io import BytesIO
import base64
from PIL import Image
import io
import json
import os
from datetime import datetime
import PyPDF2
import re
import secrets
import string

# محاولة استيراد PyMuPDF لاستخراج الصور من PDF
try:
    import fitz  # PyMuPDF for PDF rendering
    HAS_PYMUPDF = True
except Exception:
    HAS_PYMUPDF = False

app = Flask(__name__)
CORS(app)

# مسارات الملفات (مبنية على متغير البيئة DATA_DIR للتخزين الدائم)
DATA_DIR = os.environ.get("DATA_DIR", os.path.dirname(os.path.abspath(__file__)))
# مسار ملف المدربين
TRAINERS_FILE = os.path.join(DATA_DIR, 'trainers.json')
# مسار ملف النماذج المحفوظة
FORMS_FILE = os.path.join(DATA_DIR, 'saved_forms.json')

def load_trainers():
    """تحميل بيانات المدربين من الملف"""
    if os.path.exists(TRAINERS_FILE):
        with open(TRAINERS_FILE, 'r', encoding='utf-8') as f:
            return json.load(f)
    return {"trainers": []}

def save_trainers(data):
    """حفظ بيانات المدربين إلى الملف"""
    with open(TRAINERS_FILE, 'w', encoding='utf-8') as f:
        json.dump(data, f, ensure_ascii=False, indent=2)

def load_saved_forms():
    """تحميل النماذج المحفوظة من الملف"""
    if os.path.exists(FORMS_FILE):
        with open(FORMS_FILE, 'r', encoding='utf-8') as f:
            return json.load(f)
    return {"forms": []}

def save_forms(data):
    """حفظ النماذج إلى الملف"""
    with open(FORMS_FILE, 'w', encoding='utf-8') as f:
        json.dump(data, f, ensure_ascii=False, indent=2)

def add_trainer(name, username, password, role='trainer'):
    """إضافة مدرب جديد"""
    data = load_trainers()
    new_id = str(len(data['trainers']) + 1)
    new_trainer = {
        "id": new_id,
        "name": name,
        "username": username,
        "password": password,
        "role": role
    }
    data['trainers'].append(new_trainer)
    save_trainers(data)
    return new_trainer

def verify_trainer(username, password):
    """التحقق من صحة بيانات تسجيل الدخول"""
    data = load_trainers()
    for trainer in data['trainers']:
        if trainer['username'] == username and trainer['password'] == password:
            return trainer
    return None

def highlight_paragraph(paragraph, color_hex):
    p = paragraph._p
    pPr = p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    pPr.append(shd)

def extract_text_from_file(file):
    """استخراج النص من ملف مرفوع"""
    try:
        filename = file.filename.lower()
        
        if filename.endswith('.pdf'):
            return extract_text_from_pdf(file)
        elif filename.endswith('.docx'):
            return extract_text_from_docx(file)
        elif filename.endswith('.txt'):
            return extract_text_from_txt(file)
        elif filename.endswith('.rtf'):
            return extract_text_from_rtf(file)
        else:
            raise ValueError("نوع الملف غير مدعوم. الملفات المدعومة: PDF, DOCX, TXT, RTF")
            
    except Exception as e:
        raise Exception(f"خطأ في قراءة الملف: {str(e)}")

def extract_text_from_pdf(file):
    """استخراج النص من ملف PDF"""
    try:
        pdf_reader = PyPDF2.PdfReader(file)
        text = ""
        for page in pdf_reader.pages:
            text += page.extract_text() + "\n"
        return text.strip()
    except Exception as e:
        raise Exception(f"خطأ في قراءة ملف PDF: {str(e)}")

def extract_text_from_docx(file):
    """استخراج النص من ملف DOCX"""
    try:
        doc = Document(file)
        text = ""
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        return text.strip()
    except Exception as e:
        raise Exception(f"خطأ في قراءة ملف DOCX: {str(e)}")

def extract_text_from_txt(file):
    """استخراج النص من ملف TXT"""
    try:
        content = file.read().decode('utf-8')
        file.seek(0)  # إعادة تعيين المؤشر
        return content.strip()
    except UnicodeDecodeError:
        try:
            file.seek(0)
            content = file.read().decode('latin-1')
            file.seek(0)
            return content.strip()
        except Exception as e:
            raise Exception(f"خطأ في قراءة ملف TXT: {str(e)}")

def extract_text_from_rtf(file):
    """استخراج النص من ملف RTF"""
    try:
        content = file.read().decode('utf-8')
        file.seek(0)
        # إزالة تنسيق RTF
        text = re.sub(r'\\[a-z0-9-]+\d?', '', content)
        text = re.sub(r'\{[^}]*\}', '', text)
        text = re.sub(r'\\\'[0-9a-f]{2}', '', text)
        text = re.sub(r'\\\n', '\n', text)
        text = re.sub(r'\\\r', '\r', text)
        text = re.sub(r'\\\t', '\t', text)
        text = re.sub(r'\\\s', ' ', text)
        return text.strip()
    except Exception as e:
        raise Exception(f"خطأ في قراءة ملف RTF: {str(e)}")

def clean_extracted_text(text):
    """تنظيف النص المستخرج"""
    if not text:
        return ""
    
    # إزالة الأسطر الفارغة المتكررة
    text = re.sub(r'\n\s*\n', '\n\n', text)
    
    # إزالة المسافات الزائدة
    text = re.sub(r' +', ' ', text)
    
    # تنظيف النص
    text = text.strip()
    
    return text

def extract_images_from_pdf_bytes(file_bytes, max_pages=5, dpi=144):
    """استخراج الصور من ملف PDF كـ bytes"""
    if not HAS_PYMUPDF:
        return []
    
    try:
        import io
        from PIL import Image
        
        # فتح PDF من bytes
        pdf_document = fitz.open(stream=file_bytes, filetype="pdf")
        images = []
        
        # تحديد عدد الصفحات للفحص
        pages_to_check = min(max_pages, len(pdf_document))
        
        for page_num in range(pages_to_check):
            page = pdf_document[page_num]
            
            # استخراج الصور من الصفحة
            image_list = page.get_images()
            
            for img_index, img in enumerate(image_list):
                try:
                    # الحصول على بيانات الصورة
                    xref = img[0]
                    base_image = pdf_document.extract_image(xref)
                    image_bytes = base_image["image"]
                    
                    # تحويل إلى PIL Image للتحكم في الجودة
                    pil_image = Image.open(io.BytesIO(image_bytes))
                    
                    # تحويل إلى RGB إذا كان ضروري
                    if pil_image.mode in ('RGBA', 'LA', 'P'):
                        pil_image = pil_image.convert('RGB')
                    
                    # تغيير الحجم إذا كان كبير جداً
                    max_dimension = 800
                    if max(pil_image.size) > max_dimension:
                        ratio = max_dimension / max(pil_image.size)
                        new_size = tuple(int(dim * ratio) for dim in pil_image.size)
                        pil_image = pil_image.resize(new_size, Image.Resampling.LANCZOS)
                    
                    # تحويل إلى bytes مع ضغط JPEG
                    output_buffer = io.BytesIO()
                    pil_image.save(output_buffer, format='JPEG', quality=85, optimize=True)
                    compressed_image_bytes = output_buffer.getvalue()
                    
                    # تحويل إلى base64
                    import base64
                    image_b64 = base64.b64encode(compressed_image_bytes).decode('utf-8')
                    
                    images.append({
                        "mime_type": "image/jpeg",
                        "data": image_b64
                    })
                    
                    # إيقاف إذا وصلنا للحد الأقصى
                    if len(images) >= 10:
                        break
                        
                except Exception as img_error:
                    print(f"⚠️ تعذر استخراج صورة {img_index} من الصفحة {page_num}: {img_error}")
                    continue
            
            # إيقاف إذا وصلنا للحد الأقصى
            if len(images) >= 10:
                break
        
        pdf_document.close()
        print(f"✅ تم استخراج {len(images)} صورة من PDF")
        return images
        
    except Exception as e:
        print(f"❌ خطأ في استخراج الصور من PDF: {str(e)}")
        return []

def extract_images_from_docx_bytes(file_bytes, max_images=10):
    """استخراج الصور من ملف DOCX كـ bytes"""
    try:
        import io
        from PIL import Image
        
        # فتح DOCX من bytes
        doc = Document(io.BytesIO(file_bytes))
        images = []
        
        # البحث عن الصور في العلاقات
        for rel in doc.part.rels.values():
            if "image" in rel.target_ref:
                try:
                    # الحصول على بيانات الصورة
                    image_bytes = rel.target_part.blob
                    
                    # تحويل إلى PIL Image
                    pil_image = Image.open(io.BytesIO(image_bytes))
                    
                    # تحويل إلى RGB إذا كان ضروري
                    if pil_image.mode in ('RGBA', 'LA', 'P'):
                        pil_image = pil_image.convert('RGB')
                    
                    # تغيير الحجم إذا كان كبير جداً
                    max_dimension = 800
                    if max(pil_image.size) > max_dimension:
                        ratio = max_dimension / max(pil_image.size)
                        new_size = tuple(int(dim * ratio) for dim in pil_image.size)
                        pil_image = pil_image.resize(new_size, Image.Resampling.LANCZOS)
                    
                    # تحويل إلى bytes مع ضغط JPEG
                    output_buffer = io.BytesIO()
                    pil_image.save(output_buffer, format='JPEG', quality=85, optimize=True)
                    compressed_image_bytes = output_buffer.getvalue()
                    
                    # تحويل إلى base64
                    import base64
                    image_b64 = base64.b64encode(compressed_image_bytes).decode('utf-8')
                    
                    images.append({
                        "mime_type": "image/jpeg",
                        "data": image_b64
                    })
                    
                    # إيقاف إذا وصلنا للحد الأقصى
                    if len(images) >= max_images:
                        break
                        
                except Exception as img_error:
                    print(f"⚠️ تعذر استخراج صورة من DOCX: {img_error}")
                    continue
        
        print(f"✅ تم استخراج {len(images)} صورة من DOCX")
        return images
        
    except Exception as e:
        print(f"❌ خطأ في استخراج الصور من DOCX: {str(e)}")
        return []

def safe_extract_text_from_pdf_bytes(file_bytes):
    """استخراج النص من PDF بطريقة آمنة"""
    try:
        if HAS_PYMUPDF:
            # استخدام PyMuPDF إذا كان متاح
            pdf_document = fitz.open(stream=file_bytes, filetype="pdf")
            text = ""
            
            for page_num in range(len(pdf_document)):
                page = pdf_document[page_num]
                text += page.get_text() + "\n"
            
            pdf_document.close()
            return text.strip()
        else:
            # استخدام PyPDF2 كبديل
            import io
            pdf_reader = PyPDF2.PdfReader(io.BytesIO(file_bytes))
            text = ""
            
            for page in pdf_reader.pages:
                text += page.extract_text() + "\n"
            
            return text.strip()
            
    except Exception as e:
        print(f"⚠️ تعذر استخراج النص من PDF: {str(e)}")
        return ""

def safe_extract_text_from_docx_bytes(file_bytes):
    """استخراج النص من DOCX بطريقة آمنة"""
    try:
        import io
        doc = Document(io.BytesIO(file_bytes))
        text = ""
        
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        
        return text.strip()
        
    except Exception as e:
        print(f"⚠️ تعذر استخراج النص من DOCX: {str(e)}")
        return ""

@app.route('/login', methods=['POST'])
def login():
    """نقطة نهائية لتسجيل الدخول"""
    data = request.json
    username = data.get('username', '').strip()
    password = data.get('password', '').strip()
    
    if not username or not password:
        return jsonify({'success': False, 'message': 'يرجى إدخال اسم المستخدم وكلمة المرور'}), 400
    
    trainer = verify_trainer(username, password)
    if trainer:
        # تحديد الدور (افتراضي مدرب، والمشرف العام هو حساب باسم "نورا صباغ")
        role = trainer.get('role', 'trainer')
        if trainer.get('name') == 'نورا صباغ':
            role = 'admin'
            trainer['role'] = 'admin'
        return jsonify({
            'success': True, 
            'trainer': {
                'id': trainer['id'],
                'name': trainer['name'],
                'username': trainer['username'],
                'role': role
            }
        })
    else:
        return jsonify({'success': False, 'message': 'اسم المستخدم أو كلمة المرور غير صحيحة'}), 401

@app.route('/add-trainer', methods=['POST'])
def add_new_trainer():
    """إضافة مدرب جديد - متاح للمشرف فقط عبر الواجهة"""
    data = request.json
    requester_id = (data or {}).get('requesterId')
    name = (data or {}).get('name', '').strip()
    username = (data or {}).get('username', '').strip()
    password = (data or {}).get('password', '').strip()

    # التحقق من تسجيل الدخول
    if not requester_id:
        return jsonify({'success': False, 'message': 'يرجى تسجيل الدخول أولاً'}), 401

    trainers_data = load_trainers()
    requester = next((t for t in trainers_data.get('trainers', []) if t.get('id') == requester_id), None)
    if not requester or requester.get('role') != 'admin':
        return jsonify({'success': False, 'message': 'غير مصرح لك بإضافة مدربين'}), 403

    # التحقق من الحقول المطلوبة
    if not name or not username or not password:
        return jsonify({'success': False, 'message': 'جميع الحقول مطلوبة'}), 400

    # التحقق من عدم وجود اسم مستخدم مكرر
    for trainer in trainers_data['trainers']:
        if trainer.get('username') == username:
            return jsonify({'success': False, 'message': 'اسم المستخدم موجود مسبقاً'}), 400

    # دائماً اجعل الدور "trainer" للحسابات الجديدة من واجهة المشرف
    role = 'trainer'

    new_trainer = add_trainer(name, username, password, role)
    return jsonify({
        'success': True,
        'trainer': {
            'id': new_trainer['id'],
            'name': new_trainer['name'],
            'username': new_trainer['username'],
            'role': new_trainer.get('role', 'trainer')
        }
    })

@app.route('/save-form', methods=['POST'])
def save_form():
    """حفظ نموذج الأسئلة"""
    data = request.json
    trainer_id = data.get('trainerId')
    lesson_name = data.get('lessonName', '').strip()
    trainer_name = data.get('trainerName', '').strip()
    exam_name = data.get('examName', '').strip()
    questions = data.get('questions', [])
    # حقول اختيارية للتوثيق عند قيام المشرف بإنشاء نسخة
    original_form_id = data.get('originalFormId')
    original_trainer_id = data.get('originalTrainerId')
    modified_by_admin = bool(data.get('modifiedByAdmin', False))
    
    if not trainer_id:
        return jsonify({'success': False, 'message': 'يرجى تسجيل الدخول أولاً'}), 401
    
    if not lesson_name or not trainer_name:
        return jsonify({'success': False, 'message': 'يرجى إدخال اسم الدرس واسم المدرب'}), 400
    
    if not questions:
        return jsonify({'success': False, 'message': 'يرجى إضافة أسئلة على الأقل'}), 400
    
    # تحميل النماذج المحفوظة
    forms_data = load_saved_forms()
    trainers_data = load_trainers()
    trainer_obj = next((t for t in trainers_data.get('trainers', []) if t.get('id') == trainer_id), None)
    trainer_username = trainer_obj.get('username') if trainer_obj else ''
    
    # إنشاء معرف فريد للنموذج
    form_id = str(len(forms_data['forms']) + 1)
    
    # إنشاء النموذج الجديد
    new_form = {
        'id': form_id,
        'trainerId': trainer_id,
        'trainerUsername': trainer_username,
        'lessonName': lesson_name,
        'trainerName': trainer_name,
        'examName': exam_name,
        'questions': questions,
        'createdAt': datetime.now().isoformat(),
        'updatedAt': datetime.now().isoformat(),
        'lastModifiedById': trainer_id,
        'lastModifiedByName': trainer_obj['name'] if trainer_obj else trainer_name,
        'lastModifiedByUsername': trainer_username,
        'lastModifiedByRole': (trainer_obj.get('role') if trainer_obj else 'trainer')
    }

    # إضافة بيانات التوثيق إذا توفرت
    if original_form_id:
        new_form['originalFormId'] = original_form_id
    if original_trainer_id:
        new_form['originalTrainerId'] = original_trainer_id
    if modified_by_admin:
        new_form['modifiedByAdmin'] = True
    
    forms_data['forms'].append(new_form)
    save_forms(forms_data)
    
    return jsonify({
        'success': True,
        'message': 'تم حفظ النموذج بنجاح',
        'formId': form_id
    })

@app.route('/get-forms/<trainer_id>', methods=['GET'])
def get_trainer_forms(trainer_id):
    """الحصول على نماذج المدرب أو كل النماذج إذا كان المستخدم مشرفاً عاماً"""
    if not trainer_id:
        return jsonify({'success': False, 'message': 'معرف المدرب مطلوب'}), 400

    # دعم تمرير باراميتر ?role=admin من الواجهة الأمامية لتمييز المشرف
    role = request.args.get('role', 'trainer')

    forms_data = load_saved_forms()

    if role == 'admin':
        # المشرف يرى كل النماذج مع إغناء ببيانات حساب المدرب (username)
        trainers_data = load_trainers()
        trainer_id_to_username = {t.get('id'): t.get('username') for t in trainers_data.get('trainers', [])}
        for form in forms_data['forms']:
            if not form.get('trainerUsername'):
                form['trainerUsername'] = trainer_id_to_username.get(form.get('trainerId'), '')
        return jsonify({'success': True, 'forms': forms_data['forms']})

    # غير ذلك: أرجع نماذج المدرب فقط
    trainer_forms = [form for form in forms_data['forms'] if form['trainerId'] == trainer_id]
    return jsonify({'success': True, 'forms': trainer_forms})

@app.route('/get-form/<form_id>', methods=['GET'])
def get_form(form_id):
    """الحصول على نموذج محدد"""
    if not form_id:
        return jsonify({'success': False, 'message': 'معرف النموذج مطلوب'}), 400
    
    forms_data = load_saved_forms()
    form = next((f for f in forms_data['forms'] if f['id'] == form_id), None)
    
    if not form:
        return jsonify({'success': False, 'message': 'النموذج غير موجود'}), 404
    
    return jsonify({
        'success': True,
        'form': form
    })

@app.route('/update-form/<form_id>', methods=['PUT'])
def update_form(form_id):
    """تحديث نموذج موجود"""
    data = request.json
    trainer_id = data.get('trainerId')
    lesson_name = data.get('lessonName', '').strip()
    trainer_name = data.get('trainerName', '').strip()
    exam_name = data.get('examName', '').strip()
    questions = data.get('questions', [])
    
    if not trainer_id:
        return jsonify({'success': False, 'message': 'يرجى تسجيل الدخول أولاً'}), 401
    
    if not lesson_name or not trainer_name:
        return jsonify({'success': False, 'message': 'يرجى إدخال اسم الدرس واسم المدرب'}), 400
    
    if not questions:
        return jsonify({'success': False, 'message': 'يرجى إضافة أسئلة على الأقل'}), 400
    
    forms_data = load_saved_forms()
    trainers_data = load_trainers()
    trainer_obj = next((t for t in trainers_data.get('trainers', []) if t.get('id') == trainer_id), None)
    trainer_username = trainer_obj.get('username') if trainer_obj else ''
    form = next((f for f in forms_data['forms'] if f['id'] == form_id), None)
    
    if not form:
        return jsonify({'success': False, 'message': 'النموذج غير موجود'}), 404
    
    # السماح للمشرف العام بالتعديل على أي نموذج
    is_admin = trainer_obj and trainer_obj.get('role') == 'admin'
    if form['trainerId'] != trainer_id and not is_admin:
        return jsonify({'success': False, 'message': 'غير مصرح لك بتعديل هذا النموذج'}), 403
    
    # تحديث النموذج
    form['lessonName'] = lesson_name
    form['trainerName'] = trainer_name
    # تأكيد وجود اسم مستخدم المدرب داخل السجل (لدعم المشرف)
    if not form.get('trainerUsername'):
        form['trainerUsername'] = trainer_username
    form['examName'] = exam_name
    form['questions'] = questions
    form['updatedAt'] = datetime.now().isoformat()
    # تتبع آخر معدّل
    form['lastModifiedById'] = trainer_id
    form['lastModifiedByName'] = trainer_obj['name'] if trainer_obj else trainer_name
    form['lastModifiedByUsername'] = trainer_username
    form['lastModifiedByRole'] = trainer_obj.get('role') if trainer_obj else 'trainer'
    
    save_forms(forms_data)
    
    return jsonify({
        'success': True,
        'message': 'تم تحديث النموذج بنجاح'
    })

@app.route('/delete-form/<form_id>', methods=['DELETE'])
def delete_form(form_id):
    """حذف نموذج"""
    data = request.json
    trainer_id = data.get('trainerId')
    
    if not trainer_id:
        return jsonify({'success': False, 'message': 'يرجى تسجيل الدخول أولاً'}), 401
    
    forms_data = load_saved_forms()
    form = next((f for f in forms_data['forms'] if f['id'] == form_id), None)
    
    if not form:
        return jsonify({'success': False, 'message': 'النموذج غير موجود'}), 404
    
    # السماح للمشرف العام بحذف أي نموذج
    trainers_data = load_trainers()
    trainer_obj = next((t for t in trainers_data.get('trainers', []) if t.get('id') == trainer_id), None)
    is_admin = trainer_obj and trainer_obj.get('role') == 'admin'
    if form['trainerId'] != trainer_id and not is_admin:
        return jsonify({'success': False, 'message': 'غير مصرح لك بحذف هذا النموذج'}), 403
    
    # حذف النموذج
    forms_data['forms'] = [f for f in forms_data['forms'] if f['id'] != form_id]
    save_forms(forms_data)
    
    return jsonify({
        'success': True,
        'message': 'تم حذف النموذج بنجاح'
    })

@app.route('/preview-form', methods=['POST'])
def preview_form():
    """معاينة النموذج كـ HTML"""
    data = request.json
    lesson_name = data.get('lessonName', '').strip()
    trainer_name = data.get('trainerName', '').strip()
    exam_name = data.get('examName', '').strip()
    questions = data.get('questions', [])
    
    if not lesson_name or not trainer_name:
        return jsonify({'success': False, 'message': 'يرجى إدخال اسم الدرس واسم المدرب'}), 400
    
    if not questions:
        return jsonify({'success': False, 'message': 'يرجى إضافة أسئلة على الأقل'}), 400
    
    # إنشاء HTML للمعاينة
    html_content = generate_preview_html(lesson_name, trainer_name, exam_name, questions)
    
    return jsonify({
        'success': True,
        'html': html_content
    })

def generate_preview_html(lesson_name, trainer_name, exam_name, questions):
    """إنشاء HTML للمعاينة"""
    letters = ['أ', 'ب', 'ج', 'د', 'ه', 'و']
    
    # إنشاء العنوان مع اسم الاختبار
    title = lesson_name
    if exam_name:
        title += f" - {exam_name}"
    title += f" ({trainer_name})"
    
    html = f"""
    <!DOCTYPE html>
    <html lang="ar" dir="rtl">
    <head>
        <meta charset="UTF-8">
        <meta name="viewport" content="width=device-width, initial-scale=1.0">
        <title>{title}</title>
        <style>
            @import url('https://fonts.googleapis.com/css2?family=Cairo:wght@400;500;600;700&display=swap');
            body {{
                font-family: 'Cairo', sans-serif;
                line-height: 1.6;
                margin: 40px;
                background: white;
                color: #333;
            }}
            .header {{
                text-align: center;
                margin-bottom: 30px;
                padding-bottom: 20px;
                border-bottom: 2px solid #4f46e5;
            }}
            .header h1 {{
                color: #4f46e5;
                margin: 0;
                font-size: 24px;
                font-weight: 700;
            }}
            .question {{
                margin-bottom: 25px;
                padding: 15px;
                border: 1px solid #e5e7eb;
                border-radius: 8px;
                background: #f9fafb;
            }}
            .question-text {{
                font-weight: 600;
                margin-bottom: 15px;
                color: #374151;
            }}
            .question-images {{
                display: flex;
                flex-wrap: wrap;
                gap: 10px;
                margin: 10px 0;
            }}
            .question-image {{
                max-width: 300px;
                max-height: 200px;
                border-radius: 8px;
                box-shadow: 0 2px 4px rgba(0,0,0,0.1);
            }}
            .options {{
                margin-right: 20px;
            }}
            .option {{
                margin: 8px 0;
                padding: 8px 12px;
                background: white;
                border-radius: 6px;
                border: 1px solid #d1d5db;
            }}
            .correct-answer {{
                background: #dcfce7 !important;
                border-color: #22c55e !important;
                color: #166534;
                font-weight: 600;
            }}
            .classic-answer {{
                background: #f3f4f6;
                padding: 15px;
                border-radius: 8px;
                border-right: 4px solid #4f46e5;
                margin-top: 10px;
            }}
        </style>
    </head>
    <body>
        <div class="header">
            <h1>{title}</h1>
        </div>
    """
    
    for idx, q in enumerate(questions, 1):
        html += f'<div class="question">'
        html += f'<div class="question-text">{idx}. {q["text"]}</div>'
        
        # إضافة الصور إذا وجدت
        images = q.get('images', [])
        if images and len(images) > 0:
            html += '<div class="question-images">'
            for img_src in images:
                html += f'<img src="{img_src}" class="question-image" alt="صورة السؤال">'
            html += '</div>'
        
        if q['type'] in ['true_false', 'multiple_choice']:
            html += '<div class="options">'
            options = ['صح', 'خطأ'] if q['type'] == 'true_false' else q['options']
            
            for i, opt in enumerate(options):
                is_correct = opt.strip() == q['answer'].strip()
                option_class = 'option correct-answer' if is_correct else 'option'
                html += f'<div class="{option_class}">{letters[i]}. {opt}</div>'
            
            html += '</div>'
        
        elif q['type'] == 'classic':
            html += f'<div class="classic-answer"><strong>الإجابة:</strong> {q["answer"] or "..."}</div>'
        
        html += '</div>'
    
    html += '</body></html>'
    return html

@app.route('/export', methods=['POST'])
def export_questions():
    """نقطة نهائية لتصدير الأسئلة مع التحقق من تسجيل الدخول"""
    data = request.json
    trainer_id = data.get('trainerId')
    lesson_name = data.get('lessonName', '').strip()
    trainer_name = data.get('trainerName', '').strip()
    exam_name = data.get('examName', '').strip()
    questions = data.get('questions', [])

    # التحقق من وجود معرف المدرب
    if not trainer_id:
        return jsonify({'success': False, 'message': 'يرجى تسجيل الدخول أولاً'}), 401

    doc = Document()
    
    # إنشاء العنوان مع اسم الاختبار
    title = lesson_name
    if exam_name:
        title += f" - {exam_name}"
    title += f" ({trainer_name})"
    
    doc.add_heading(title, 0)
    doc.add_paragraph("")

    letters = ['a', 'b', 'c', 'd', 'e', 'f']

    for idx, q in enumerate(questions, 1):
        doc.add_paragraph(f"{idx}. {q['text']}")
        # إضافة الصور إذا وجدت
        images = q.get('images', [])
        if images and len(images) > 0:
            for img_src in images:
                try:
                    img_data = img_src.split(',')[1] if ',' in img_src else img_src
                    img_bytes = base64.b64decode(img_data)
                    img_stream = io.BytesIO(img_bytes)
                    doc.add_picture(img_stream, width=Inches(2.5))
                except Exception as e:
                    doc.add_paragraph("[تعذر إضافة الصورة]")
                    print(f"Error adding image for question {idx}: {e}")
                    # متابعة العملية حتى لو حدث خطأ في إضافة الصورة

        if q['type'] in ['true_false', 'multiple_choice']:
            options = ['صح', 'خطأ'] if q['type'] == 'true_false' else q['options']
            for i, opt in enumerate(options):
                # فقرة من درجة ثانية (إزاحة لليمين)
                p = doc.add_paragraph(f"{letters[i]}. {opt}")
                p.paragraph_format.left_indent = Inches(0.4)  # درجة ثانية
                if opt.strip() == q['answer'].strip():
                    highlight_paragraph(p, '38761D')  # أخضر غامق

        elif q['type'] == 'classic':
            p = doc.add_paragraph("الإجابة: ...")
            p.paragraph_format.left_indent = Inches(0.4)

        doc.add_paragraph("")  # سطر فارغ بين الأسئلة

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    filename = f"{title}.docx"
    return send_file(buffer, as_attachment=True, download_name=filename)

@app.route('/test', methods=['GET'])
def test_connection():
    """نقطة اختبار بسيطة للاتصال"""
    return jsonify({
        'success': True,
        'message': 'الخادم يعمل بشكل صحيح',
        'status': 'running',
        'timestamp': datetime.now().isoformat()
    })

@app.route('/upload-lesson-file', methods=['POST'])
def upload_lesson_file():
    """رفع ملف درس واستخراج النص منه"""
    try:
        if 'file' not in request.files:
            return jsonify({'success': False, 'message': 'لم يتم اختيار ملف'}), 400
        
        file = request.files['file']
        if file.filename == '':
            return jsonify({'success': False, 'message': 'لم يتم اختيار ملف'}), 400
        
        # التحقق من نوع الملف
        allowed_extensions = {'.pdf', '.docx', '.txt', '.rtf', '.jpg', '.jpeg', '.png', '.gif', '.bmp', '.webp'}
        file_ext = os.path.splitext(file.filename.lower())[1]
        
        if file_ext not in allowed_extensions:
            return jsonify({'success': False, 'message': f'نوع الملف غير مدعوم. الملفات المدعومة: {", ".join(allowed_extensions)}'}), 400
        
        # إرجاع نجاح فوري لجميع أنواع الملفات - سيتم تحليلها مباشرة في Gemini
        return jsonify({
            'success': True,
            'extractedText': '',
            'textLength': 0,
            'fileName': file.filename,
            'message': f'تم تحميل الملف "{file.filename}" بنجاح. سيتم تحليله مباشرة في Gemini عند توليد الأسئلة.'
        })
        
    except Exception as e:
        print(f"خطأ في upload_lesson_file: {str(e)}")
        return jsonify({'success': False, 'message': f'خطأ في معالجة الملف: {str(e)}'}), 500

@app.route('/suggest-questions', methods=['POST'])
def suggest_questions():
    """اقتراح أسئلة باستخدام Google Gemini API الجديد بناءً على الملف مباشرة"""
    try:
        # تهيئة المتغيرات
        lesson_file = None
        lesson_text = None
        use_file_directly = False
        file_data = None
        
        # التحقق من وجود ملف أو نص
        if 'file' in request.files:
            # إذا كان هناك ملف، استخدمه مباشرة
            file = request.files['file']
            print(f"📁 تم استلام ملف: {file.filename}")
            
            if file.filename == '':
                return jsonify({'success': False, 'message': 'لم يتم اختيار ملف'}), 400
            
            # التحقق من نوع الملف
            allowed_extensions = {'.pdf', '.docx', '.txt', '.rtf', '.jpg', '.jpeg', '.png', '.gif', '.bmp'}
            file_ext = os.path.splitext(file.filename.lower())[1]
            
            if file_ext not in allowed_extensions:
                return jsonify({'success': False, 'message': f'نوع الملف غير مدعوم. الملفات المدعومة: {", ".join(allowed_extensions)}'}), 400
            
            # استخدام الملف مباشرة - قراءة المحتوى فوراً
            try:
                file_content = file.read()
                print(f"📁 تم قراءة الملف بنجاح، الحجم: {len(file_content)} بايت")
                
                # تحديد نوع MIME
                mime_type = "application/octet-stream"
                filename = file.filename.lower()
                if filename.endswith('.pdf'):
                    mime_type = "application/pdf"
                elif filename.endswith('.docx'):
                    mime_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                elif filename.endswith('.txt'):
                    mime_type = "text/plain"
                elif filename.endswith('.rtf'):
                    mime_type = "application/rtf"
                elif filename.endswith(('.jpg', '.jpeg')):
                    mime_type = "image/jpeg"
                elif filename.endswith('.png'):
                    mime_type = "image/png"
                elif filename.endswith('.gif'):
                    mime_type = "image/gif"
                elif filename.endswith('.bmp'):
                    mime_type = "image/bmp"
                
                print(f"📁 نوع MIME المحدد: {mime_type}")
                
                # تحويل إلى base64
                file_content_b64 = base64.b64encode(file_content).decode('utf-8')
                print(f"📁 تم تحويل الملف إلى base64، الحجم: {len(file_content_b64)} حرف")
                
                use_file_directly = True
                file_data = {
                    'content': file_content_b64,
                    'mime_type': mime_type,
                    'filename': file.filename
                }
                print(f"📁 استخدام الملف مباشرة: {file.filename}")
                
            except Exception as file_error:
                print(f"❌ خطأ في قراءة الملف: {str(file_error)}")
                return jsonify({'success': False, 'message': f'خطأ في قراءة الملف: {str(file_error)}'}), 500
            
        elif 'lessonText' in request.form or 'lessonText' in request.json:
            # إذا كان هناك نص، استخدمه (للتوافق مع الإصدارات السابقة)
            if request.is_json:
                data = request.json
                lesson_text = data.get('lessonText', '').strip()
            else:
                lesson_text = request.form.get('lessonText', '').strip()
            
            lesson_file = None
            use_file_directly = False
            print(f"📝 استخدام النص المقدم: {len(lesson_text)} حرف")
            
        else:
            return jsonify({'success': False, 'message': 'يجب تقديم ملف أو نص الدرس'}), 400
        
        # استخراج باقي البيانات
        if request.is_json:
            data = request.json
        else:
            data = request.form
        
        # معالجة question_count - قد يكون string أو رقم
        question_count_raw = data.get('questionCount', 5)
        try:
            question_count = int(question_count_raw)
            print(f"📊 question_count: {question_count}")
        except (ValueError, TypeError) as e:
            print(f"❌ خطأ في تحويل question_count: {str(e)}")
            question_count = 5
        
        # معالجة question_types - قد يكون string أو dict
        question_types_raw = data.get('questionTypes', {})
        if isinstance(question_types_raw, str):
            try:
                question_types = json.loads(question_types_raw)
                print(f"📝 تم تحويل question_types من string إلى dict: {question_types}")
            except json.JSONDecodeError as e:
                print(f"❌ خطأ في تحليل question_types: {str(e)}")
                question_types = {}
        else:
            question_types = question_types_raw
            
        detailed_request = data.get('detailedRequest', False)  # هل الطلب مفصل أم لا
        if isinstance(detailed_request, str):
            detailed_request = detailed_request.lower() == 'true'
        
        # عداد المحاولات
        max_retries = 3
        current_retry = 0
        
        # التحقق من وجود محتوى صالح
        has_valid_content = False
        if lesson_text and lesson_text.strip():
            has_valid_content = True
        elif lesson_file:
            has_valid_content = True
        elif use_file_directly and file_data:
            has_valid_content = True
            
        if not has_valid_content:
            return jsonify({'success': False, 'message': 'نص الدرس أو الملف مطلوب'}), 400
        
        print(f"🔍 بدء اقتراح الأسئلة...")
        print(f"📊 إجمالي عدد الأسئلة المطلوبة: {question_count}")
        
        if detailed_request and question_types and isinstance(question_types, dict):
            print(f"🎯 طلب مفصل:")
            print(f"   - صح وخطأ: {question_types.get('true_false', 0)}")
            print(f"   - متعدد الخيارات: {question_types.get('multiple_choice', 0)}")
            print(f"   - كتابي: {question_types.get('classic', 0)}")
        else:
            print(f"📋 طلب بسيط: {question_count} أسئلة متنوعة")
            # تعيين قيم افتراضية إذا لم تكن هناك أنواع محددة
            if not question_types or not isinstance(question_types, dict):
                question_types = {
                    'true_false': max(1, question_count // 3),
                    'multiple_choice': max(1, question_count // 2),
                    'classic': max(1, question_count // 6)
                }
                print(f"📝 تم تعيين قيم افتراضية: {question_types}")
        
        # مفتاح API
        api_key = os.environ.get("GOOGLE_API_KEY", "AIzaSyD_bgcfFpGr82pCAcO1MH9zCvppicQje5o")
        
        # URL الجديد لـ Gemini API
        url = "https://generativelanguage.googleapis.com/v1beta/models/gemini-2.0-flash:generateContent"
        
        # إنشاء prompt محسن بناءً على نوع الطلب
        if detailed_request and question_types and isinstance(question_types, dict):
            # طلب مفصل مع أعداد محددة لكل نوع
            true_false_count = question_types.get('true_false', 0)
            multiple_choice_count = question_types.get('multiple_choice', 0)
            classic_count = question_types.get('classic', 0)
            
            prompt = f"""
            بناءً على المحتوى التالي (قد يتضمن نص وصور ورسوم بيانية وجداول)، اقترح أسئلة تعليمية بالعدد المحدد لكل نوع بالضبط:

            المطلوب بالضبط (يجب الالتزام بهذه الأعداد تماماً):
            - أسئلة صح/خطأ (true_false): {true_false_count} سؤال
            - أسئلة متعددة الخيارات (multiple_choice): {multiple_choice_count} سؤال  
            - أسئلة كتابية (classic): {classic_count} سؤال

            المجموع الإجمالي المطلوب: {true_false_count + multiple_choice_count + classic_count} سؤال

            أعد الإجابة بتنسيق JSON فقط بالشكل التالي:
            {{
                "questions": [
                    {{
                        "type": "multiple_choice",
                        "text": "نص السؤال",
                        "options": ["الخيار الأول", "الخيار الثاني", "الخيار الثالث", "الخيار الرابع"],
                        "answer": "الخيار الصحيح"
                    }},
                    {{
                        "type": "true_false", 
                        "text": "نص السؤال",
                        "answer": "صح"
                    }},
                    {{
                        "type": "classic",
                        "text": "نص السؤال",
                        "answer": "الإجابة النموذجية"
                    }}
                ]
            }}

            تعليمات مهمة جداً:
            1. يجب أن يكون العدد الإجمالي للأسئلة {true_false_count + multiple_choice_count + classic_count} بالضبط
            2. يجب أن يكون عدد أسئلة صح/خطأ {true_false_count} بالضبط
            3. يجب أن يكون عدد أسئلة متعددة الخيارات {multiple_choice_count} بالضبط  
            4. يجب أن يكون عدد الأسئلة الكتابية {classic_count} بالضبط
            5. لا تزيد ولا تنقص عن هذه الأعداد
            6. تأكد من أن الأسئلة مناسبة للمستوى التعليمي
            7. تأكد من أن الأسئلة تغطي محتوى الملف بشكل شامل (بما في ذلك الصور والرسوم البيانية)
            8. تأكد من أن الأسئلة واضحة ومفهومة
            9. تأكد من أن الأسئلة متنوعة في الصعوبة
            10. إذا كان هناك صور أو رسوم بيانية، تأكد من إنشاء أسئلة عنها
            """
        else:
            # طلب بسيط (للتوافق مع الإصدارات السابقة)
            prompt = f"""
            بناءً على المحتوى التالي (قد يتضمن نص وصور ورسوم بيانية وجداول)، اقترح {question_count} أسئلة تعليمية متنوعة.
    
            المطلوب:
            1. أسئلة متعددة الخيارات (multiple_choice)
            2. أسئلة صح/خطأ (true_false) 
            3. أسئلة كتابية (classic)
            
            أعد الإجابة بتنسيق JSON فقط بالشكل التالي:
            {{
                "questions": [
                    {{
                        "type": "multiple_choice",
                        "text": "نص السؤال",
                        "options": ["الخيار الأول", "الخيار الثاني", "الخيار الثالث", "الخيار الرابع"],
                        "answer": "الخيار الصحيح"
                    }},
                    {{
                        "type": "true_false", 
                        "text": "نص السؤال",
                        "answer": "صح"
                    }},
                    {{
                        "type": "classic",
                        "text": "نص السؤال",
                        "answer": "الإجابة النموذجية"
                    }}
                ]
            }}
            
            تأكد من أن الأسئلة:
            - مناسبة للمستوى التعليمي
            - تغطي محتوى الملف بشكل شامل (بما في ذلك الصور والرسوم البيانية)
            - واضحة ومفهومة
            - متنوعة في الصعوبة
            - إذا كان هناك صور أو رسوم بيانية، تأكد من إنشاء أسئلة عنها
            """
        
        # حلقة المحاولات
        while current_retry < max_retries:
            current_retry += 1
            print(f"📤 المحاولة {current_retry}/{max_retries} - إرسال طلب إلى Gemini 2.0 Flash...")
            
            try:
                # إعداد بيانات الطلب بناءً على نوع المحتوى
                if use_file_directly:
                    # استخدام الملف مباشرة
                    print(f"📁 استخدام بيانات الملف المحفوظة: {file_data['filename']}")
                    
                    # بناء محتوى متعدد الوسائط للطلب
                    parts = []
                    parts.append({"text": prompt})
                    
                    added_any_image = False
                    
                    try:
                        _mime = file_data['mime_type']
                        _b64 = file_data['content']
                        _bytes = base64.b64decode(_b64)
                        
                        if _mime.startswith("image/"):
                            # صورة مباشرة - إرسالها كما هي
                            parts.append({"inline_data": {"mime_type": _mime, "data": _b64}})
                            added_any_image = True
                            print(f"📁 إرسال صورة مباشرة: {file_data['filename']}")
                            
                        elif _mime == "application/pdf":
                            # PDF: إرسال الملف مباشرة + استخراج الصور كدعم إضافي
                            parts.append({"inline_data": {"mime_type": _mime, "data": _b64}})
                            added_any_image = True
                            
                            # استخراج الصور كدعم إضافي (اختياري)
                            imgs = extract_images_from_pdf_bytes(_bytes, max_pages=3, dpi=144)
                            if imgs:
                                print(f"📁 تم استخراج {len(imgs)} صورة إضافية من PDF كدعم")
                                for img in imgs:
                                    parts.append({"inline_data": img})
                            
                            print(f"📁 إرسال PDF مباشرة: {file_data['filename']}")
                            
                        elif _mime in ("application/vnd.openxmlformats-officedocument.wordprocessingml.document", "application/msword"):
                            # DOCX: إرسال الملف مباشرة + استخراج الصور كدعم إضافي
                            parts.append({"inline_data": {"mime_type": _mime, "data": _b64}})
                            added_any_image = True
                            
                            # استخراج الصور كدعم إضافي (اختياري)
                            imgs = extract_images_from_docx_bytes(_bytes, max_images=5)
                            if imgs:
                                print(f"📁 تم استخراج {len(imgs)} صورة إضافية من DOCX كدعم")
                                for img in imgs:
                                    parts.append({"inline_data": img})
                            
                            print(f"📁 إرسال DOCX مباشرة: {file_data['filename']}")
                            
                        else:
                            # أنواع أخرى (TXT/RTF): إرسال الملف مباشرة
                            parts.append({"inline_data": {"mime_type": _mime, "data": _b64}})
                            added_any_image = True
                            print(f"📁 إرسال ملف نصي مباشرة: {file_data['filename']}")
                                
                    except Exception as _e:
                        print(f"⚠️ تعذر تجهيز الملف مباشرة: {_e}")
                        # في حالة الفشل، نرسل الملف كما هو
                        parts.append({"inline_data": {"mime_type": file_data['mime_type'], "data": file_data['content']}})
                        added_any_image = True
                    
                    request_data = {"contents": [{"parts": parts}]}
                    
                    print(f"📁 إرسال الملف مباشرة: {file_data['filename']} (نوع: {file_data['mime_type']})")
                    print(f"📁 عدد الأجزاء المرسلة: {len(parts)}")
                    
                else:
                    # استخدام النص
                    request_data = {
                        "contents": [
                            {
                                "parts": [
                                    {
                                        "text": prompt + f"\n\nنص الدرس:\n{lesson_text}"
                                    }
                                ]
                            }
                        ]
                    }
                    
                    print(f"📝 إرسال النص: {len(lesson_text)} حرف")
                
                # رؤوس الطلب
                headers = {
                    'Content-Type': 'application/json',
                    'X-goog-api-key': api_key
                }
                
                # إرسال الطلب
                print(f"📤 إرسال طلب إلى Gemini API...")
                print(f"📤 URL: {url}")
                print(f"📤 Headers: {headers}")
                print(f"📤 Request data keys: {list(request_data.keys()) if isinstance(request_data, dict) else 'FormData'}")
                
                try:
                    response = requests.post(
                        url,
                        headers=headers,
                        json=request_data,
                        timeout=120  # زيادة الوقت للملفات الكبيرة
                    )
                    
                    print(f"📥 استجابة API: {response.status_code}")
                    
                    if response.status_code != 200:
                        print(f"❌ خطأ في API: {response.status_code}")
                        print(f"📝 تفاصيل الخطأ: {response.text}")
                        if current_retry >= max_retries:
                            return jsonify({'success': False, 'message': f'خطأ في API: {response.status_code} - {response.text}'}), 500
                        continue
                        
                except requests.exceptions.RequestException as req_error:
                    print(f"❌ خطأ في طلب HTTP: {str(req_error)}")
                    if current_retry >= max_retries:
                        return jsonify({'success': False, 'message': f'خطأ في الاتصال: {str(req_error)}'}), 500
                    continue
                
                # تحليل الاستجابة
                response_data = response.json()
                print(f"📝 بيانات الاستجابة: {str(response_data)[:200]}...")
        
                # استخراج النص من الاستجابة
                if 'candidates' in response_data and len(response_data['candidates']) > 0:
                    candidate = response_data['candidates'][0]
                    if 'content' in candidate and 'parts' in candidate['content']:
                        parts = candidate['content']['parts']
                        if len(parts) > 0 and 'text' in parts[0]:
                            response_text = parts[0]['text'].strip()
                            print(f"✅ تم استلام النص من Gemini")
                            print(f"📝 النص المستلم: {response_text[:200]}...")
                        else:
                            print("❌ لم يتم العثور على نص في الاستجابة")
                            if current_retry >= max_retries:
                                return jsonify({'success': False, 'message': 'لم يتم العثور على نص في الاستجابة'}), 500
                            continue
                    else:
                        print("❌ تنسيق الاستجابة غير صحيح")
                        if current_retry >= max_retries:
                            return jsonify({'success': False, 'message': 'تنسيق الاستجابة غير صحيح'}), 500
                        continue
                else:
                    print("❌ لم يتم العثور على مرشح في الاستجابة")
                    if current_retry >= max_retries:
                        return jsonify({'success': False, 'message': 'لم يتم العثور على مرشح في الاستجابة'}), 500
                    continue
                
                # محاولة استخراج JSON من النص
                try:
                    # البحث عن JSON في النص
                    start_idx = response_text.find('{')
                    end_idx = response_text.rfind('}') + 1
                    
                    if start_idx != -1 and end_idx != 0:
                        json_text = response_text[start_idx:end_idx]
                        print(f"🔍 JSON المستخرج: {json_text[:200]}...")
                
                        questions_data = json.loads(json_text)
                
                        # التحقق من صحة البيانات
                        if 'questions' in questions_data and isinstance(questions_data['questions'], list):
                            questions = questions_data['questions']
                            print(f"✅ تم العثور على {len(questions)} سؤال")
                            
                            # التحقق من عدد الأسئلة لكل نوع
                            if detailed_request and question_types and isinstance(question_types, dict):
                                type_counts = {}
                                for q in questions:
                                    q_type = q.get('type', 'unknown')
                                    type_counts[q_type] = type_counts.get(q_type, 0) + 1
                                
                                print(f"📊 توزيع الأسئلة المستلمة:")
                                print(f"   - صح وخطأ: {type_counts.get('true_false', 0)}")
                                print(f"   - متعدد الخيارات: {type_counts.get('multiple_choice', 0)}")
                                print(f"   - كتابي: {type_counts.get('classic', 0)}")
                                
                                # التحقق من تطابق الأعداد
                                expected_true_false = question_types.get('true_false', 0)
                                expected_multiple_choice = question_types.get('multiple_choice', 0)
                                expected_classic = question_types.get('classic', 0)
                                
                                actual_true_false = type_counts.get('true_false', 0)
                                actual_multiple_choice = type_counts.get('multiple_choice', 0)
                                actual_classic = type_counts.get('classic', 0)
                                
                                total_expected = expected_true_false + expected_multiple_choice + expected_classic
                                total_actual = len(questions)
                                
                                if (actual_true_false != expected_true_false or 
                                    actual_multiple_choice != expected_multiple_choice or 
                                    actual_classic != expected_classic or 
                                    total_actual != total_expected):
                                    
                                    error_msg = f"""
                                    المحاولة {current_retry}/{max_retries} - عدد الأسئلة المستلمة لا يتطابق مع المطلوب:
                                    
                                    المطلوب:
                                    - صح وخطأ: {expected_true_false}
                                    - متعدد الخيارات: {expected_multiple_choice}
                                    - كتابي: {expected_classic}
                                    - المجموع: {total_expected}
                                    
                                    المستلم:
                                    - صح وخطأ: {actual_true_false}
                                    - متعدد الخيارات: {actual_multiple_choice}
                                    - كتابي: {actual_classic}
                                    - المجموع: {total_actual}
                                    """
                                    
                                    print(f"❌ {error_msg}")
                                    
                                    # إذا كانت هذه آخر محاولة، أعد رسالة خطأ
                                    if current_retry >= max_retries:
                                        final_error_msg = error_msg + "\n\nتم استنفاذ جميع المحاولات. يرجى المحاولة مرة أخرى."
                                        return jsonify({'success': False, 'message': final_error_msg}), 400
                                    
                                    # إذا لم تكن آخر محاولة، استمر في الحلقة
                                    print(f"🔄 المحاولة مرة أخرى... ({current_retry + 1}/{max_retries})")
                                    continue  # الانتقال إلى المحاولة التالية
                                
                                print(f"✅ تم التحقق من تطابق الأعداد")
                                # إذا وصلنا هنا، نجحت المحاولة
                                return jsonify({
                                    'success': True,
                                    'questions': questions,
                                    'message': f'تم اقتراح {len(questions)} سؤال بنجاح'
                                })
                            
                            # إذا لم تكن هناك حاجة للتحقق المفصل، أعد النتيجة
                            return jsonify({
                                'success': True,
                                'questions': questions,
                                'message': f'تم اقتراح {len(questions)} سؤال بنجاح'
                            })
                        else:
                            print("❌ تنسيق البيانات غير صحيح")
                            if current_retry >= max_retries:
                                return jsonify({'success': False, 'message': 'تنسيق البيانات غير صحيح'}), 500
                            continue
                    else:
                        print("❌ لم يتم العثور على JSON في الاستجابة")
                        if current_retry >= max_retries:
                            return jsonify({'success': False, 'message': 'لم يتم العثور على بيانات صحيحة'}), 500
                        continue
                
                except json.JSONDecodeError as e:
                    print(f"❌ خطأ في تحليل JSON: {str(e)}")
                    print(f"📝 النص الذي فشل في تحليله: {response_text}")
                    if current_retry >= max_retries:
                        return jsonify({'success': False, 'message': f'خطأ في تحليل البيانات: {str(e)}'}), 500
                    continue
                    
            except Exception as e:
                print(f"❌ خطأ في المحاولة {current_retry}: {str(e)}")
                if current_retry >= max_retries:
                    return jsonify({'success': False, 'message': f'حدث خطأ أثناء المحاولة: {str(e)}'}), 500
                continue
        
        # إذا وصلنا هنا، فشلت جميع المحاولات
        return jsonify({'success': False, 'message': 'فشلت جميع المحاولات. يرجى المحاولة مرة أخرى.'}), 500
            
    except Exception as e:
        print(f"❌ خطأ في suggest_questions: {str(e)}")
        print(f"🔍 نوع الخطأ: {type(e).__name__}")
        import traceback
        traceback.print_exc()
        return jsonify({'success': False, 'message': f'حدث خطأ أثناء اقتراح الأسئلة: {str(e)}'}), 500

@app.route('/suggest-trainer-credentials', methods=['POST'])
def suggest_trainer_credentials():
    """اقتراح اسم مستخدم وكلمة مرور قوية عند إنشاء مدرب جديد"""
    data = request.json or {}
    name = data.get('name', '').strip()

    # توليد اسم مستخدم بناءً على الاسم أو بشكل عشوائي
    trainers_data = load_trainers()
    existing_usernames = set(t['username'] for t in trainers_data.get('trainers', []))

    def generate_username(base_name):
        base = re.sub(r'[^a-zA-Z0-9]', '', base_name.lower()) if base_name else 'trainer'
        for i in range(100):
            suffix = secrets.token_hex(2)[:3]  # 3 أحرف عشوائية
            username = f"{base}{suffix}"
            if username not in existing_usernames:
                return username
        # fallback
        while True:
            username = 'trainer' + secrets.token_hex(3)
            if username not in existing_usernames:
                return username

    username = generate_username(name.split()[0] if name else '')

    # توليد كلمة مرور قوية
    def generate_password(length=12):
        alphabet = string.ascii_letters + string.digits + string.punctuation
        while True:
            password = ''.join(secrets.choice(alphabet) for _ in range(length))
            # تحقق من القوة: على الأقل حرف كبير وصغير ورقم ورمز
            if (any(c.islower() for c in password) and any(c.isupper() for c in password)
                and any(c.isdigit() for c in password) and any(c in string.punctuation for c in password)):
                return password

    password = generate_password()

    return jsonify({
        'success': True,
        'username': username,
        'password': password
    })

@app.route('/all-trainers', methods=['GET'])
def get_all_trainers():
    """إرجاع قائمة جميع المدربين (اسم، اسم مستخدم، كلمة مرور، الرتبة)"""
    data = load_trainers()
    trainers = [
        {
            'name': t.get('name', '') if t.get('name') is not None else '',
            'username': t.get('username', '') if t.get('username') is not None else '',
            'password': t.get('password', '') if t.get('password') is not None else '',
            'role': t.get('role', 'trainer')
        }
        for t in data.get('trainers', [])
    ]
    return jsonify({'success': True, 'trainers': trainers})

@app.route('/edit-trainer', methods=['POST'])
def edit_trainer():
    data = request.json or {}
    old_username = data.get('old_username', '').strip()
    name = data.get('name', '').strip()
    username = data.get('username', '').strip()
    password = data.get('password', '').strip()
    if not old_username or not name or not username or not password:
        return jsonify({'success': False, 'message': 'جميع الحقول مطلوبة'}), 400
    trainers_data = load_trainers()
    trainers = trainers_data.get('trainers', [])
    trainer = next((t for t in trainers if t.get('username') == old_username), None)
    if not trainer:
        return jsonify({'success': False, 'message': 'المدرب غير موجود'}), 404
    # تحقق من عدم تكرار اسم المستخدم الجديد (إلا إذا لم يتغير)
    if username != old_username and any(t.get('username') == username for t in trainers):
        return jsonify({'success': False, 'message': 'اسم المستخدم موجود مسبقاً'}), 400
    trainer['name'] = name
    trainer['username'] = username
    trainer['password'] = password
    save_trainers(trainers_data)
    return jsonify({'success': True})

if __name__ == '__main__':
    port = int(os.environ.get("PORT", 5000))
    app.run(host="0.0.0.0", port=port)
